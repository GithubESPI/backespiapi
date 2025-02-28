import base64
import json
import shutil
import time
import zipfile  # Assurez-vous d'importer le module zipfile standard

# Importation des modules nécessaires
import asyncio
import fitz  # PyMuPDF
from fastapi import FastAPI, HTTPException, APIRouter, WebSocket, WebSocketDisconnect
from fastapi.responses import JSONResponse, FileResponse
from openpyxl import load_workbook
from app.core.config import settings
import openpyxl
from docx import Document
import os
import logging
from pydantic import BaseModel
import requests
import subprocess

from app.services.api_service import fetch_api_data
from app.services.excel_service import process_excel_file, update_excel_with_appreciations
from app.utils.date_utils import sum_durations
from starlette.websockets import WebSocketDisconnect
from datetime import datetime



# Configuration du logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Création de l'application FastAPI
app = FastAPI()

# Création d'un routeur pour organiser les routes
router = APIRouter()

# Variable globale pour stocker la progression
# Liste des connexions WebSocket actives
active_connections = []
progress_data = {}

# Définition du modèle de réponse pour les uploads
class UploadResponse(BaseModel):
    configId: str

# Définition du modèle pour les URLs des documents
class DocumentUrls(BaseModel):
    sessionId: str
    excelUrl: str
    wordUrl: str


# Fonction pour normaliser les titres en supprimant les caractères non alphanumériques et en les mettant en minuscules
def normalize_title(title):
    import re
    if not isinstance(title, str):
        title = str(title)
    return re.sub(r'\W+', '', title).lower()


class_ids = {
    "MAPI": "ID_MAPI_001",
    "MAGI": "ID_MAGI_001",
    "MEFIM": "ID_MEFIM_001",
    "MAPI_S2": "ID_MAPI_S2_001",
    "MAGI_S2": "ID_MAGI_S2_001",
    "MEFIM_S2": "ID_MEFIM_S2_001",
    "MAPI_S3": "ID_MAPI_S3_001",
    "MAGI_S3": "ID_MAGI_S3_001",
    "MEFIM_S3": "ID_MEFIM_S3_001",
    "MAPI_S4": "ID_MAPI_S4_001",
    "MAGI_S4": "ID_MAGI_S4_001",
    "MEFIM_S4": "ID_MEFIM_S4_001",
    "BG_ALT_S1": "ID_BG_ALT_S1_001",
    "BG_ALT_S2": "ID_BG_ALT_S2_001",
    "BG_ALT_S3": "ID_BG_ALT_S3_001",
    "BG_ALT_S4": "ID_BG_ALT_S4_001",
    "BG_ALT_S5": "ID_BG_ALT_S5_001",
    "BG_ALT_S6": "ID_BG_ALT_S6_001",
    "BG_TP_1": "ID_BG_TP_1_001",
    "BG_TP_2": "ID_BG_TP_2_001",
    "BG_TP_3": "ID_BG_TP_3_001",
    "BG_TP_4": "ID_BG_TP_4_001",
    "BG_TP_5": "ID_BG_TP_5_001",
    "BG_TP_6": "ID_BG_TP_6_001"
}

# Fonction pour récupérer les données d'API en parallèle
async def fetch_api_data_for_template(headers, class_name=None):
    api_url_mapping = {
        "MAGI": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "MAGI_S2": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "MAGI_S3": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "MAGI_S4": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "MEFIM": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "MEFIM_S2": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "MEFIM_S3": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "MEFIM_S4": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "MAPI": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "MAPI_S2": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "MAPI_S3": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "MAPI_S4": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "BG_ALT_S1": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "BG_ALT_S2": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "BG_ALT_S3": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "BG_ALT_S4": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "BG_ALT_S5": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "BG_ALT_S6": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "BG_TP_1": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "BG_TP_2": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "BG_TP_3": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "BG_TP_4": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "BG_TP_5": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        },
        "BG_TP_6": {
            "apprenants_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            "groupes_url": f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2"
        }
    }


    if class_name and class_name in api_url_mapping:
        urls = api_url_mapping[class_name]
        api_urls = [
            urls["apprenants_url"],
            urls["groupes_url"],
            f"https://groupe-espi.ymag.cloud/index.php/r/v1/absences/01-09-2023/15-09-2024",
            f"https://groupe-espi.ymag.cloud/index.php/r/v1/apprenants/frequentes?codesPeriode=2",
            f"https://groupe-espi.ymag.cloud/index.php/r/v1/periodes"
        ]
    else:
        api_urls = [
            f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/apprenants?codesPeriode=2",
            f"https://groupe-espi.ymag.cloud/index.php/r/v1/formation-longue/groupes?codesPeriode=2",
            f"https://groupe-espi.ymag.cloud/index.php/r/v1/absences/01-09-2023/15-09-2024",
            f"https://groupe-espi.ymag.cloud/index.php/r/v1/apprenants/frequentes?codesPeriode=2",
            f"https://groupe-espi.ymag.cloud/index.php/r/v1/periodes"
        ]

    api_data_futures = [fetch_api_data(url, headers) for url in api_urls]
    results = await asyncio.gather(*api_data_futures, return_exceptions=True)

    for i, result in enumerate(results):
        if isinstance(result, Exception):
            logger.error(f"API request failed for {api_urls[i]}: {result}")
        elif isinstance(result, dict):
            logger.debug(f"API response from {api_urls[i]}: {result}")
        else:
            logger.error(f"Unexpected response from {api_urls[i]}: {type(result)}")

    if any(isinstance(result, Exception) for result in results):
        raise HTTPException(status_code=500, detail="Failed to fetch API data")

    api_data, raw_groupes_data, absences_data, frequentes_data, periodes_dict = results
    
    # Construction du dictionnaire des groupes
    groupes_dict = {}
    for groupe in raw_groupes_data.values():
        if isinstance(groupe, dict) and 'codeGroupe' in groupe:
            groupe_code = str(groupe['codeGroupe'])  # Convertir en string pour la cohérence
            groupes_dict[groupe_code] = {
                'codeGroupe': groupe_code,
                'nomGroupe': groupe.get('nomGroupe', 'N/A'),
                'etenduGroupe': groupe.get('etenduGroupe', 'N/A')
            }
    
    logger.debug(f"Groupes data structure: {groupes_dict}")

    # Traitement des périodes
    periode_2023_2024 = next(
        (periode for periode in periodes_dict.values() 
        if periode.get('codePeriode') == 2), None
    )

    # Traitement des fréquentations avec vérification des dates
    # Traitement des fréquentations avec vérification des dates
    frequentes_dict = {}
    for frequentation in frequentes_data.values():
        code_apprenant = str(frequentation.get('codeApprenant'))  # Convertir en string
        date_debut_frequentation = frequentation.get('dateDeb')
        date_fin_frequentation = frequentation.get('dateFin')
        
        # Vérifier que toutes les données nécessaires sont présentes
        if not all([code_apprenant, date_debut_frequentation, date_fin_frequentation]):
            logger.warning(f"Missing required data for frequentation: code_apprenant={code_apprenant}, "f"date_debut={date_debut_frequentation}, date_fin={date_fin_frequentation}")
            continue

        if periode_2023_2024:
            try:
                # Vérifier que les dates sont bien au format attendu
                if not isinstance(date_debut_frequentation, str) or not isinstance(date_fin_frequentation, str):
                    logger.warning(f"Invalid date format for frequentation: date_debut={date_debut_frequentation}, "f"date_fin={date_fin_frequentation}")
                    continue

                # S'assurer que periode_2023_2024 contient les bonnes données
                periode_debut_str = periode_2023_2024.get('dateDeb')
                periode_fin_str = periode_2023_2024.get('dateFin')
                
                if not periode_debut_str or not periode_fin_str:
                    logger.warning("Missing period dates in periode_2023_2024")
                    continue

                # Parser les dates
                date_debut = datetime.strptime(date_debut_frequentation, '%d/%m/%Y')
                date_fin = datetime.strptime(date_fin_frequentation, '%d/%m/%Y')
                periode_debut = datetime.strptime(periode_debut_str, '%d/%m/%Y')
                periode_fin = datetime.strptime(periode_fin_str, '%d/%m/%Y')

                if date_debut <= periode_fin and date_fin >= periode_debut:
                    # Si l'apprenant a déjà une fréquentation, comparer les dates
                    if code_apprenant in frequentes_dict:
                        existing_date_debut_str = frequentes_dict[code_apprenant].get('dateDeb')
                        if existing_date_debut_str:
                            existing_date_debut = datetime.strptime(existing_date_debut_str, '%d/%m/%Y')
                            if date_debut > existing_date_debut:
                                frequentes_dict[code_apprenant] = frequentation
                    else:
                        frequentes_dict[code_apprenant] = frequentation

                    logger.debug(f"Added/Updated frequentation for apprenant {code_apprenant} "f"with groupe {frequentation.get('codeGroupe')}")

            except ValueError as e:
                logger.error(f"Error processing dates for apprenant {code_apprenant}: {str(e)}")
                continue
            except Exception as e:
                logger.error(f"Unexpected error processing frequentation for apprenant {code_apprenant}: {str(e)}")
                continue

    # Mise à jour des informations de groupe dans api_data
    for apprenant in api_data.values():
        code_apprenant = str(apprenant.get('codeApprenant'))
        if code_apprenant in frequentes_dict:
            frequentes_info = frequentes_dict[code_apprenant]
            code_groupe = str(frequentes_info.get('codeGroupe'))
            
            if code_groupe and code_groupe in groupes_dict:
                if not apprenant.get('informationsCourantes'):
                    apprenant['informationsCourantes'] = {}
                apprenant['informationsCourantes'].update({
                    'codeGroupe': code_groupe,
                    'nomGroupe': groupes_dict[code_groupe].get('nomGroupe'),
                    'etenduGroupe': groupes_dict[code_groupe].get('etenduGroupe')
                })

    return api_data, groupes_dict, absences_data, frequentes_dict, periodes_dict

# Function to extract appreciations from Word document
def extract_appreciations_from_word(word_path):
    try:
        doc = Document(word_path)
        appreciations = {}
        for table in doc.tables:
            for row in table.rows:
                cells = row.cells
                if len(cells) >= 2:
                    name = cells[0].text.strip()
                    appreciation = cells[1].text.strip()
                    if name and appreciation:
                        appreciations[name] = appreciation
        return appreciations
    except Exception as e:
        logger.error("Failed to extract appreciations from Word document", exc_info=True)
    return {}

# Function to log data in an Excel worksheet
def log_excel_data(worksheet):
    data = []
    for row in worksheet.iter_rows(values_only=True):
        data.append(row)
    logger.debug(f"Excel data: {data}")

# Fonction pour extraire les appréciations depuis un document Word
def extract_appreciations_from_word(word_path):
    try:
        doc = Document(word_path)
        appreciations = {}
        for table in doc.tables:
            for row in table.rows:
                cells = row.cells
                if len(cells) >= 2:
                    name = cells[0].text.strip()
                    appreciation = cells[1].text.strip()
                    if name and appreciation:
                        appreciations[name] = appreciation
        return appreciations
    except Exception as e:
        logger.error("Failed to extract appreciations from Word document", exc_info=True)
    return {}

# Fonction pour logger les données d'un worksheet Excel
def log_excel_data(worksheet):
    data = []
    for row in worksheet.iter_rows(values_only=True):
        data.append(row)
    logger.debug(f"Excel data: {data}")

def extract_code_apprenant(pdf_path: str) -> str:
    try:
        with fitz.open(pdf_path) as pdf_document:
            for page_num in range(pdf_document.page_count):
                page = pdf_document.load_page(page_num)
                text = page.get_text("text")
                lines = text.split('\n')
                for line in lines:
                    if "Identifiant :" in line:
                        parts = line.split("Identifiant :")
                        if len(parts) > 1:
                            code_apprenant = parts[1].strip()
                            if code_apprenant.replace('.', '', 1).isdigit():
                                return str(int(float(code_apprenant)))
        logger.error(f"No valid code_apprenant found in {pdf_path}")
        return None
    except Exception as e:
        logger.error(f"Failed to extract code_apprenant from {pdf_path}", exc_info=True)
        return None

def import_document_to_yparéo(file_path, code_apprenant, retries=3, delay=5):
    for attempt in range(retries):
        try:
            with open(file_path, 'rb') as pdf_file:
                file_content = pdf_file.read()
                encoded_content = base64.b64encode(file_content).decode('utf-8')

            # Log the file being uploaded
            logger.info(f"Attempting to upload {file_path} for apprenant {code_apprenant}")

            # Création du JSON payload pour l'API
            payload = {
                "contenu": encoded_content,
                "nomDocument": os.path.basename(file_path),
                "typeMime": "application/pdf",
                "extension": "pdf",
            }

            # Endpoint Yparéo
            endpoint = f"/r/v1/document/apprenant/{code_apprenant}/document?codeRepertoire=1000011"
            url = f"{settings.YPAERO_BASE_URL}{endpoint}"
            headers = {
                "X-Auth-Token": settings.YPAERO_API_TOKEN,
                "Content-Type": "application/json"
            }

            # Envoi de la requête POST
            response = requests.post(url, headers=headers, json=payload)

            # Log the API response
            if response.status_code == 200:
                logger.info(f"Successfully uploaded {file_path} for apprenant {code_apprenant}")
                return True
            else:
                logging.error(f"Attempt {attempt + 1} failed with status code {response.status_code}: {response.text}")
                if response.status_code == 500:
                    raise ValueError(f"Server error while importing document {file_path}")
        
        except Exception as e:
            logging.error(f"Attempt {attempt + 1} failed due to exception: {str(e)}", exc_info=True)

        time.sleep(delay)
    
    logger.error(f"Failed to import document {file_path} after {retries} retries")
    raise ValueError(f"Server error while importing document {file_path} after {retries} retries")

# Process the uploaded file and integrate data into the Excel template
async def process_file(uploaded_wb, template_path, columns_config, class_name, current_periode, previous_periode, api_data, frequentes_dict, websocket=None):
    try:
        template_wb = openpyxl.load_workbook(template_path, data_only=True)
        uploaded_ws = uploaded_wb.active
        template_ws = template_wb.active

        header_row_uploaded = 4
        header_row_template = 1

        # Configuration des en-têtes et correspondances
        uploaded_titles = {
            normalize_title(uploaded_ws.cell(row=header_row_uploaded, column=col).value): col
            for col in range(1, uploaded_ws.max_column + 1)
            if uploaded_ws.cell(row=header_row_uploaded, column=col).value is not None
        }

        template_titles = {
            normalize_title(template_ws.cell(row=header_row_template, column=col).value): col
            for col in range(1, template_ws.max_column + 1)
            if template_ws.cell(row=header_row_template, column=col).value is not None
        }

        matching_columns = {
            uploaded_title: (uploaded_titles[uploaded_title], template_titles[template_title])
            for uploaded_title in uploaded_titles
            for template_title in template_titles
            if uploaded_title == template_title
        }

        if not matching_columns:
            return JSONResponse(content={"message": "No matching columns found."})

        # Obtention des données API
        headers = {
            'X-Auth-Token': settings.YPAERO_API_TOKEN,
            'Content-Type': 'application/json'
        }

        api_data, groupes_dict, absences_data, frequentes_dict, periodes_dict = await fetch_api_data_for_template(headers, class_name)

        if not isinstance(api_data, dict) or not isinstance(groupes_dict, dict) or not isinstance(absences_data, dict):
            raise HTTPException(status_code=500, detail="Unexpected API response format")

        # Création du dictionnaire des apprenants
        api_dict = {
            normalize_title(apprenant['nomApprenant'] + apprenant['prenomApprenant']): apprenant 
            for apprenant in api_data.values()
        }

        # Préparation du résumé des absences
        absences_summary = {}
        for absence in absences_data.values():
            apprenant_id = str(absence.get('codeApprenant'))
            if not apprenant_id:
                continue

            duration = int(absence.get('duree', 0))
            if apprenant_id not in absences_summary:
                absences_summary[apprenant_id] = {'justified': [], 'unjustified': [], 'delays': []}

            if absence.get('isJustifie'):
                absences_summary[apprenant_id]['justified'].append(duration)
            elif absence.get('isRetard'):
                absences_summary[apprenant_id]['delays'].append(duration)
            else:
                absences_summary[apprenant_id]['unjustified'].append(duration)

        # Traitement des lignes
        exclude_phrase = 'moyennedugroupe'
        total_rows = uploaded_ws.max_row - header_row_uploaded
        processed_rows = 0

        for row in range(header_row_uploaded + 1, uploaded_ws.max_row + 1):
            processed_rows += 1
            progress = (processed_rows / total_rows) * 100

            if websocket:
                try:
                    await websocket.send(json.dumps({
                        "progress": progress,
                        "message": f"Processing row {processed_rows} of {total_rows}"
                    }))
                except WebSocketDisconnect:
                    pass

            if any(exclude_phrase in normalize_title(uploaded_ws.cell(row=row, column=col).value or '')
                for col in range(1, uploaded_ws.max_column + 1)):
                continue

            uploaded_name = uploaded_ws.cell(row=row, column=columns_config['name_column_index_uploaded']).value
            if not uploaded_name:
                continue

            template_row = row - header_row_uploaded + header_row_template + 1
            template_ws.cell(row=template_row, column=columns_config['name_column_index_template']).value = uploaded_name

            normalized_name = normalize_title(uploaded_name)
            if apprenant_info := api_dict.get(normalized_name):
                # Remplissage des informations de base
                template_ws.cell(row=template_row, column=columns_config['code_apprenant_column_index_template']).value = apprenant_info.get('codeApprenant', 'N/A')
                template_ws.cell(row=template_row, column=columns_config['date_naissance_column_index_template']).value = apprenant_info.get('dateNaissance', 'N/A')

                # Informations du site
                if 'inscriptions' in apprenant_info and apprenant_info['inscriptions']:
                    template_ws.cell(row=template_row, column=columns_config['nom_site_column_index_template']).value = apprenant_info['inscriptions'][0]['site'].get('nomSite', 'N/A')

                # Informations du groupe
                code_groupe = str(apprenant_info.get('informationsCourantes', {}).get('codeGroupe'))
                if code_groupe in groupes_dict:
                    groupe_info = groupes_dict[code_groupe]
                    template_ws.cell(row=template_row, column=columns_config['code_groupe_column_index_template']).value = groupe_info.get('codeGroupe', 'N/A')
                    template_ws.cell(row=template_row, column=columns_config['nom_groupe_column_index_template']).value = groupe_info.get('nomGroupe', 'N/A')
                    template_ws.cell(row=template_row, column=columns_config['etendu_groupe_column_index_template']).value = groupe_info.get('etenduGroupe', 'N/A')

                # Informations des absences
                apprenant_id = str(apprenant_info.get('codeApprenant'))
                if apprenant_id in absences_summary:
                    abs_info = absences_summary[apprenant_id]
                    template_ws.cell(row=template_row, column=columns_config['duree_justifie_column_index_template']).value = sum_durations(abs_info['justified']) or "00h00"
                    template_ws.cell(row=template_row, column=columns_config['duree_non_justifie_column_index_template']).value = sum_durations(abs_info['unjustified']) or "00h00"
                    template_ws.cell(row=template_row, column=columns_config['duree_retard_column_index_template']).value = sum_durations(abs_info['delays']) or "00h00"

            # Copie des autres colonnes correspondantes
            for uploaded_title, (src_col, dest_col) in matching_columns.items():
                template_ws.cell(row=template_row, column=dest_col).value = uploaded_ws.cell(row=row, column=src_col).value

        # Nettoyage final du template
        for col in range(1, template_ws.max_column + 1):
            if template_ws.cell(row=header_row_template + 1, column=col).value == template_ws.cell(row=header_row_template, column=col).value:
                template_ws.cell(row=header_row_template + 1, column=col).value = None
            if template_ws.cell(row=header_row_template + 2, column=col).value == "Note":
                template_ws.cell(row=header_row_template + 2, column=col).value = None

        return template_wb

    except Exception as e:
        logger.error("Failed to process the file", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))
    
def process_periodes_data(periodes_dict):
    current_periode = None
    previous_periode = None
    
    for periode in periodes_dict.values():
        date_debut = periode.get('dateDeb')
        if date_debut is None:
            logger.warning(f"Missing 'dateDeb' in periode: {periode}")
            continue

        if current_periode is None or date_debut > current_periode.get('dateDeb', ''):
            current_periode = periode
    
    for periode in periodes_dict.values():
        date_debut = periode.get('dateDeb')
        if date_debut is None:
            continue

        if date_debut < current_periode.get('dateDeb', ''):
            if previous_periode is None or date_debut > previous_periode.get('dateDeb', ''):
                previous_periode = periode
    
    if current_periode is None:
        logger.error("No valid periode found with 'dateDeb'")
    
    return current_periode, previous_periode

def convert_docx_to_pdf(docx_dir):
    libreoffice_path = 'soffice' # Remplacez par le chemin correct de LibreOffice

    for filename in os.listdir(docx_dir):
        if filename.endswith('.docx'):
            docx_path = os.path.join(docx_dir, filename)
            pdf_path = os.path.join(docx_dir, filename.replace('.docx', '.pdf'))

            command = [libreoffice_path, '--headless', '--convert-to', 'pdf', '--outdir', docx_dir, docx_path]

            try:
                subprocess.run(command, check=True)
                if os.path.exists(pdf_path):
                    logger.info(f"Converted {docx_path} to {pdf_path}")
                else:
                    logger.error(f"PDF not created for: {docx_path}")
            except subprocess.CalledProcessError as e:
                logger.error(f"Failed to convert {docx_path} to PDF: {e}")
            except FileNotFoundError as e:
                logger.error(f"LibreOffice executable not found: {e}")

def clean_output_directory(output_dir):
    try:
        if os.path.exists(output_dir):
            shutil.rmtree(output_dir)
        os.makedirs(output_dir, exist_ok=True)
        logger.info(f"Cleaned and recreated the output directory: {output_dir}")
    except Exception as e:
        logger.error(f"Failed to clean output directory: {output_dir}", exc_info=True)

# Fonction pour déduire le nom de la classe à partir du fichier Excel
# Fonction pour déduire le nom de la classe à partir du fichier Excel
def determine_class_name(uploaded_values, current_periode):
    logger.debug(f"Uploaded values for class name detection: {uploaded_values}")

    class_name = None
    class_id = None
    periode_code = current_periode.get('codePeriode') if current_periode else None

    # Check for specific course names and map to class names accordingly for MAGI and MEFIM
    if any("Baux Commerciaux et Gestion Locative" in value for value in uploaded_values):
        class_name = "MAGI"
    elif any("Rénovation Energétique des Actifs Tertiaires" in value for value in uploaded_values):
        class_name = "MAGI_S3"
    elif any("Business Game Property Management" in value for value in uploaded_values):
        class_name = "MAGI_S4" 
    elif any("Budget d'Exploitation et de Travaux" in value for value in uploaded_values):
        class_name = "MAGI_S2"
    
    # Logic for MEFIM
    elif any("Les Fondamentaux de l'Evaluation" in value for value in uploaded_values):
        class_name = "MEFIM"
    elif any("Marché d'Actifs Immobiliers" in value for value in uploaded_values):
        class_name = "MEFIM_S2"
    elif any("Droit des Suretés et de la Transmission" in value for value in uploaded_values):
        class_name = "MEFIM_S3"
    elif any("Business Game Arbitrage et Stratégies d'Investissement" in value for value in uploaded_values):
        class_name = "MEFIM_S4"
    
    # Logic for MAPI
    elif any("Étude Foncière" in value for value in uploaded_values):
        class_name = "MAPI"
    elif any("Droit de la Promotion Immobilière" in value for value in uploaded_values):
        class_name = "MAPI_S2"
    elif any("Acquisition et Dissociation du Foncier" in value for value in uploaded_values):
        class_name = "MAPI_S3"
    elif any("Business Game Aménagement et Promotion Immobilière" in value for value in uploaded_values):
        class_name = "MAPI_S4"

    # Logic for BG_ALT_S1 to BG_ALT_S6
    elif any("Économie Générale" in value for value in uploaded_values):
        class_name = "BG_ALT_S1"
    elif any("Microéconomie I" in value for value in uploaded_values):
        class_name = "BG_ALT_S2"
    elif any("Microéconomie II" in value for value in uploaded_values):
        class_name = "BG_ALT_S3"
    elif any("Marketing Digital & Environnemental" in value for value in uploaded_values):
        class_name = "BG_ALT_S4"
    elif any("Économie Urbaine" in value for value in uploaded_values):
        class_name = "BG_ALT_S5"
    elif any("Finance Immobilière" in value for value in uploaded_values):
        class_name = "BG_ALT_S6"

    # Logic for BG_TP_1 to BG_TP_6
    elif any("Organisations, Stratégies et Innovations I" in value for value in uploaded_values):
        class_name = "BG_TP_1"
    elif any("Real Estate English & TOEFL" in value for value in uploaded_values):
        class_name = "BG_TP_2"
    elif any("Pratique de Gestion Locative I" in value for value in uploaded_values):
        class_name = "BG_TP_3"
    elif any("Mobilité Internationale Études" in value for value in uploaded_values):
        class_name = "BG_TP_4"
    elif any("Management de Projet Immobilier" in value for value in uploaded_values):
        class_name = "BG_TP_5"
    elif any("Mémoire de Recherche" in value for value in uploaded_values):
        class_name = "BG_TP_6"

    if class_name:
        class_id = class_ids.get(class_name)
        logger.debug(f"Detected class name: {class_name} with ID: {class_id} for period: {periode_code}")
        return class_name, class_id, periode_code
    else:
        logger.error("Failed to detect class name")
        return None, None, periode_code

async def update_progress(session_id: str, progress: int):
    # Mise à jour de la variable globale
    progress_data[session_id] = progress

    # Préparation du message à envoyer
    message = {
        "session_id": session_id,
        "progress": progress,
        "message": f"Processing progress: {progress}%"
    }

    # Envoi de la mise à jour à tous les clients WebSocket connectés
    for connection in active_connections:
        try:
            await connection.send_json(message)
        except WebSocketDisconnect:
            # Si la connexion est fermée, nous la retirons de la liste des connexions actives
            active_connections.remove(connection)
        except Exception as e:
            # Log any other errors that might occur during sending
            print(f"Error sending WebSocket update: {str(e)}")

# Fonction pour gérer les nouvelles connexions WebSocket
@router.websocket("/ws/progress/{session_id}")
async def websocket_endpoint(websocket: WebSocket, session_id: str):
    await websocket.accept()
    active_connections.append(websocket)
    try:
        while True:
            data = await websocket.receive_text()
            # Vous pouvez traiter les messages reçus ici si nécessaire
    except WebSocketDisconnect:
        logger.info(f"WebSocket connection closed for session {session_id}")
    finally:
        if websocket in active_connections:
            active_connections.remove(websocket)



@router.post("/upload-and-integrate-excel-and-word")
async def upload_and_integrate(doc_urls: DocumentUrls):
    global progress_data
    session_id = doc_urls.sessionId
    progress_data[session_id] = 0
    
    # Initialisation de la progression
    try:
        progress_data[session_id] = 5  # Progression à 5%
        await update_progress(session_id, 5)
        
        excel_response = requests.get(doc_urls.excelUrl)
        if excel_response.status_code != 200:
            logger.error(f"Failed to download Excel document: {excel_response.status_code}")
            raise HTTPException(status_code=400, detail="Failed to download Excel document")

        # Sauvegarder le fichier Excel temporairement
        os.makedirs(settings.DOWNLOAD_DIR, exist_ok=True)
        temp_excel_path = os.path.join(settings.DOWNLOAD_DIR, f"{doc_urls.sessionId}.xlsx")
        with open(temp_excel_path, 'wb') as temp_excel_file:
            temp_excel_file.write(excel_response.content)
            
        progress_data[session_id] = 15  # Progression à 10%
        await update_progress(session_id, 15)

        # Télécharger le fichier Word
        word_response = requests.get(doc_urls.wordUrl)
        if word_response.status_code != 200:
            raise HTTPException(status_code=400, detail="Failed to download Word document")

        temp_word_path = os.path.join(settings.DOWNLOAD_DIR, f"{doc_urls.sessionId}.docx")
        with open(temp_word_path, 'wb') as temp_word_file:
            temp_word_file.write(word_response.content)
            
        if not os.path.exists(temp_word_path):
            logger.error(f"Word document not found at {temp_word_path}")
            raise HTTPException(status_code=400, detail="Word document not found")

        # Traitement du fichier Excel
        progress_data[session_id] = 20  # Progression à 50%
        await update_progress(session_id, 20)  # Utiliser une fonction séparée pour mettre à jour la progression
        
        uploaded_wb = load_workbook(temp_excel_path, data_only=True)
        uploaded_ws = uploaded_wb.active
        
        # Ajout de progressions intermédiaires
        progress_data[session_id] = 25
        await update_progress(session_id, 25)
        
        headers = {
            'X-Auth-Token': settings.YPAERO_API_TOKEN,
            'Content-Type': 'application/json'
        }

        api_data, groupes_data, absences_data, frequentes_dict, periodes_dict = await fetch_api_data_for_template(headers, None)

        # Vérification plus souple des données
        if not all(isinstance(data, dict) for data in [api_data, groupes_data, absences_data, frequentes_dict, periodes_dict]):
            logger.error(f"Unexpected API response format. Received types: api_data: {type(api_data)}, groupes_data: {type(groupes_data)}, absences_data: {type(absences_data)}, frequentes_dict: {type(frequentes_dict)}, periodes_dict: {type(periodes_dict)}")
            raise HTTPException(status_code=500, detail="Unexpected API response format")

        # Ajoutez des logs pour inspecter les données
        logger.debug(f"API data received: {api_data}")
        logger.debug(f"Groupes data received: {groupes_data}")
        logger.debug(f"Absences data received: {absences_data}")
        logger.debug(f"Frequentes dict received: {frequentes_dict}")
        logger.debug(f"Periodes dict received: {periodes_dict}")

        # Traiter les données de période avant de les utiliser
        current_periode, previous_periode = process_periodes_data(periodes_dict)
        logger.debug(f"Current periode: {current_periode}")
        logger.debug(f"Previous periode: {previous_periode}")


        # Détection du template approprié en fonction des valeurs du fichier Excel
        uploaded_values = [uploaded_ws[cell].value for cell in ['C4', 'F4', 'I4', 'L4', 'O4', 'R4', 'U4', 'X4', 'AA4', 'AD4', 'AG4', 'AJ4', 'AM4', 'AP4', 'AS4', 'AV4', 'AY4', 'BB4', 'BE4', 'BH4', 'BK4', 'BN4', 'BQ4', 'BT4', 'BW4', 'BZ4'] if uploaded_ws[cell].value is not None]

        class_name, class_id, periode_code = determine_class_name(uploaded_values, current_periode)
        
        if not isinstance(api_data, dict) or not isinstance(groupes_data, dict) or not isinstance(absences_data, dict) or not isinstance(periodes_dict, dict):
            raise HTTPException(status_code=500, detail="Unexpected API response format")

        templates = {
            "MAPI": settings.M1_S1_MAPI_TEMPLATE,
            "MAGI": settings.M1_S1_MAGI_TEMPLATE,
            "MEFIM": settings.M1_S1_MEFIM_TEMPLATE,
            "MAPI_S2": settings.M1_S2_MAPI_TEMPLATE,
            "MAGI_S2": settings.M1_S2_MAGI_TEMPLATE,
            "MEFIM_S2": settings.M1_S2_MEFIM_TEMPLATE,
            "MAPI_S3": settings.M2_S3_MAPI_TEMPLATE,
            "MAGI_S3": settings.M2_S3_MAGI_TEMPLATE,
            "MEFIM_S3": settings.M2_S3_MEFIM_TEMPLATE,
            "MAPI_S4": settings.M2_S4_MAPI_TEMPLATE,
            "MAGI_S4": settings.M2_S4_MAGI_TEMPLATE,
            "MEFIM_S4": settings.M2_S4_MEFIM_TEMPLATE,
            "BG_ALT_S1": settings.BG_ALT_1_TEMPLATE,
            "BG_ALT_S2": settings.BG_ALT_2_TEMPLATE,
            "BG_ALT_S3": settings.BG_ALT_3_TEMPLATE,
            "BG_ALT_S4": settings.BG_ALT_4_TEMPLATE,
            "BG_ALT_S5": settings.BG_ALT_5_TEMPLATE,
            "BG_ALT_S6": settings.BG_ALT_6_TEMPLATE,
            "BG_TP_S1": settings.BG_TP_1_TEMPLATE,
            "BG_TP_S2": settings.BG_TP_2_TEMPLATE,
            "BG_TP_S3": settings.BG_TP_3_TEMPLATE,
            "BG_TP_S4": settings.BG_TP_4_TEMPLATE,
            "BG_TP_S5": settings.BG_TP_5_TEMPLATE,
            "BG_TP_S6": settings.BG_TP_6_TEMPLATE,
        }

        matching_values = {
            "MAPI": ['UE 1 – Economie & Gestion', 'Stratégie et Solutions Immobilières', 'Finance Immobilière', 'Économie Immobilière I', 'UE 2 – Droit', 'Droit des Affaires et des Contrats', 'UE 3 – Aménagement & Urbanisme', 'Ville et Développements Urbains', "Politique de l'Habitat", 'UE 4 – Compétences Professionnalisantes', 'Real Estate English', "Rencontres de l'Immobilier", 'ESPI Career Services', 'ESPI Inside', 'Immersion Professionnelle', 'Projet Voltaire', 'UE SPE – MAPI', 'Étude Foncière', "Montage d'une Opération de Promotion Immobilière", 'Acquisition et Dissociation du Foncier'],
            "MAGI": ['UE 1 – Economie & Gestion', 'Stratégie et Solutions Immobilières', 'Finance Immobilière', 'Économie Immobilière I', 'UE 2 – Droit', 'Droit des Affaires et des Contrats', 'UE 3 – Aménagement & Urbanisme', 'Ville et Développements Urbains', "Politique de l'Habitat", 'UE 4 – Compétences Professionnalisantes', 'Real Estate English', "Rencontres de l'Immobilier", 'ESPI Career Services', 'ESPI Inside', 'Immersion Professionnelle', 'Projet Voltaire', 'UE SPE – MAGI', 'Baux Commerciaux et Gestion Locative', 'Actifs Tertiaires en Copropriété', 'Techniques du Bâtiment'],
            "MEFIM": ['UE 1 – Economie & Gestion', 'Stratégie et Solutions Immobilières', 'Finance Immobilière', 'Économie Immobilière I', 'UE 2 – Droit', 'Droit des Affaires et des Contrats', 'UE 3 – Aménagement & Urbanisme', 'Ville et Développements Urbains', "Politique de l'Habitat", 'UE 4 – Compétences Professionnalisantes', 'Real Estate English', "Rencontres de l'Immobilier", 'ESPI Career Services', 'ESPI Inside', 'Immersion Professionnelle', 'Projet Voltaire', 'UE SPE – MEFIM', "Les Fondamentaux de l'Evaluation", 'Analyse et Financement Immobilier', 'Modélisation Financière'],
            
            "MAPI_S2": ['UE 1 – Economie & Gestion', "Marketing de l'Immobilier", 'Investissement et Financiarisation', 'Fiscalité', 'UE 2 – Droit', "Droit de l'Urbanisme et de la Construction", "Déontologie en France et à l'International", 'UE 4 – Compétences Professionnalisantes', 'Immersion Professionnelle', 'Real Estate English', 'Atelier Méthodologie de la Recherche', 'Techniques de Négociation', "Rencontres de l'Immobilier", 'ESPI Inside', 'Projet Voltaire', 'UE SPE – MAPI', "Droit de la Promotion Immobilière", "Montage d'une Opération de Logement", 'Financement des Opérations de Promotion Immobilière', "Logement Social et Accession Sociale"],
            "MAGI_S2": ['UE 1 – Economie & Gestion', "Marketing de l'Immobilier", 'Investissement et Financiarisation', 'Fiscalité', 'UE 2 – Droit', "Droit de l'Urbanisme et de la Construction", "Déontologie en France et à l'International", 'UE 4 – Compétences Professionnalisantes', 'Immersion Professionnelle', 'Real Estate English', 'Atelier Méthodologie de la Recherche', 'Techniques de Négociation', "Rencontres de l'Immobilier", 'ESPI Inside', 'Projet Voltaire', 'UE SPE – MAGI', "Budget d'Exploitation et de Travaux", 'Développement et Stratégie Commerciale', 'Technique et Conformité des Immeubles', "Gestion de l'Immobilier - Logistique et Data Center"],
            "MEFIM_S2": ['UE 1 – Economie & Gestion', "Marketing de l'Immobilier", 'Investissement et Financiarisation', 'Fiscalité', 'UE 2 – Droit', "Droit de l'Urbanisme et de la Construction", "Déontologie en France et à l'International", 'UE 4 – Compétences Professionnalisantes', 'Immersion Professionnelle', 'Real Estate English', 'Atelier Méthodologie de la Recherche', 'Techniques de Négociation', "Rencontres de l'Immobilier", 'ESPI Inside', 'Projet Voltaire', 'UE SPE – MEFIM', "Marché d'Actifs Immobiliers", "Baux Commerciaux", 'Évaluation des Actifs Résidentiels', "Audit et Gestion des Immeubles"],
            
            "MAPI_S3": ['UE 1 – Economie & Gestion', "PropTech et Innovation", 'Économie Immobilière II', 'UE 3 – Aménagement & Urbanisme', "Stratégies et Aménagement des Territoires I", "UE 4 – Compétences Professionnalisantes", 'Communication Digitale, Ecrite et Orale', 'Immersion Professionnelle', 'Real Estate English', 'Méthodologie de la Recherche', "Rencontres de l'Immobilier", 'ESPI Inside', 'UE SPE – MAPI', "Acquisition et Dissociation du Foncier", "Montage des Opérations Tertiaires", "Aménagement et Commande Publique", "Techniques du Bâtiment", "Réhabilitation et Pathologies du Bâtiment"],
            "MAGI_S3": ['UE 1 – Economie & Gestion', "PropTech et Innovation", 'Économie Immobilière II', 'UE 3 – Aménagement & Urbanisme', "Stratégies et Aménagement des Territoires I", "UE 4 – Compétences Professionnalisantes", 'Communication Digitale, Ecrite et Orale', 'Immersion Professionnelle', 'Real Estate English', 'Méthodologie de la Recherche', "Rencontres de l'Immobilier", 'ESPI Inside', 'UE SPE – MAGI', "Rénovation Energétique des Actifs Tertiaires", "Arbitrage, Optimisation et Valorisation des Actifs Tertiaires", 'Maintenance et Facility Management', "Réhabilitation et Pathologies du Bâtiment"],
            "MEFIM_S3": ['UE 1 – Economie & Gestion', "PropTech et Innovation", 'Économie Immobilière II', 'UE 3 – Aménagement & Urbanisme', "Stratégies et Aménagement des Territoires I", "UE 4 – Compétences Professionnalisantes", 'Communication Digitale, Ecrite et Orale', 'Immersion Professionnelle', 'Real Estate English', 'Méthodologie de la Recherche', "Rencontres de l'Immobilier", 'ESPI Inside', 'UE SPE – MEFIM', "Droit des Suretés et de la Transmission", 'Due Diligence', "Évaluation d'Actifs Tertiaires et Industriels", "Gestion de Patrimoine"],
            
            "MAPI_S4": ['UE 1 – Economie & Gestion', "Économie de l'Environnement", 'UE 3 – Aménagement & Urbanisme', "Normalisation, Labellisation", "Stratégies et Aménagement des Territoires II", 'UE 4 – Compétences Professionnalisantes', 'Real Estate English', 'Mémoire de Recherche', "Rencontres de l'Immobilier", 'ESPI Career Services', 'Immersion Professionnelle', 'UE SPE – MAPI', "Business Game Aménagement et Promotion Immobilière", "Fiscalité et Promotion Immobilière", "Contentieux de l'Urbanisme"],
            "MAGI_S4": ['UE 1 – Economie & Gestion', "Économie de l'Environnement", 'UE 3 – Aménagement & Urbanisme', "Normalisation, Labellisation", "Stratégies et Aménagement des Territoires II", 'UE 4 – Compétences Professionnalisantes', 'Real Estate English', 'Mémoire de Recherche', "Rencontres de l'Immobilier", 'ESPI Career Services', 'Immersion Professionnelle', 'UE SPE – MAGI', "Business Game Property Management", "Gestion des Centres Commerciaux", "Gestion de Contentieux et Recouvrement"],
            "MEFIM_S4": ['UE 1 – Economie & Gestion', "Économie de l'Environnement", 'UE 3 – Aménagement & Urbanisme', "Normalisation, Labellisation", "Stratégies et Aménagement des Territoires II", 'UE 4 – Compétences Professionnalisantes', 'Real Estate English', 'Mémoire de Recherche', "Rencontres de l'Immobilier", 'ESPI Career Services', 'Immersion Professionnelle', 'UE SPE – MEFIM', "Business Game Arbitrage et Stratégies d'Investissement", "Fiscalité du Patrimoine", "Fintech et Blockchain"],
            
            "BG_ALT_S1": ['UE 1 – Economie & Gestion', "Économie Générale", "Outils d'Analyse Economique", "Organisations, Stratégies et Innovations I", "UE 2 – Droit", "Introduction au Droit", "Droit des Contrats", "UE 3 – Aménagement & Urbanisme", "Introduction aux Méthodes d'Analyse et de Représentation Spatiale", "UE 4 – Compétences Professionnalisantes", "Immersion Professionnelle", "Real Estate English", "Panorama de l'Immobilier", "Expression Ecrite et Orale", "Gestion du Travail", "Déontologie et Ethique Professionnelle", "ESPI Career Services", "ESPI Inside"],
            "BG_ALT_S2": ['UE 1 – Economie & Gestion', "Microéconomie I", "Introduction à la Finance", "Marketing & Prospection", "Mathématiques Financières", "UE 2 – Droit", "Droit des Biens", "Droit de la Copropriété I", "Droit des Baux d'Habitation", "UE 3 – Aménagement & Urbanisme", "Histoire Urbaine et Architecture", "UE 4 – Compétences Professionnalisantes", "Immersion Professionnelle", "Real Estate English", "Panorama de l'Immobilier", "Expression Ecrite et Orale", "Gestion de Projet", "ESPI Career Services", "ESPI Inside"],
            "BG_ALT_S3": ['UE 1 – Economie & Gestion', "Microéconomie II", "Organisations, Stratégies et Innovations II", "Pratique de Gestion Locative I", "Enjeux de l'Immobilier et Solutions Digitales I", "Transactions Résidentielles", "UE 2 – Droit", "Droit de la Vente Immobilière", "Droit de la Copropriété II", "UE 3 – Aménagement & Urbanisme", "Technologie du Bâtiment", "UE 4 – Compétences Professionnalisantes", "Immersion Professionnelle", "Real Estate English", "Panorama de l'Immobilier", "Expression Ecrite et Orale", "ESPI Inside"],
            
            "BG_ALT_S4": ['UE 1 – Economie & Gestion', "Marketing Digital & Environnemental", "Enjeux de l'Immobilier et Solutions Digitales II", "Macroéconomie et Politiques Economiques", "UE 2 – Droit", "Droit du Numérique", "Droit de l'Urbanisme", "Fiscalité Générale", "UE 3 – Aménagement & Urbanisme", "Immobilier et Dynamiques Urbaines", "UE 4 – Compétences Professionnalisantes", "Immersion Professionnelle", "Real Estate English", "Panorama de l'Immobilier", "Expression Ecrite et Orale", "ESPI Inside"],
            "BG_ALT_S5": ['UE 1 – Economie & Gestion', "Économie Urbaine", "Pratique de Gestion Locative II", "Management de Projet Immobilier", "UE 2 – Droit", "Droit de la Transaction Immobilière", "Droit de l'Environnement", "Fiscalité Immobilière", "UE 3 – Aménagement & Urbanisme", "Habitat et Développement Durable", "UE 4 – Compétences Professionnalisantes", "Immersion Professionnelle", "Real Estate English", "Panorama de l'Immobilier", "Expression Ecrite et Orale", "Atelier Urbain I", "Méthodologie de la Recherche", "ESPI Inside"],
            "BG_ALT_S6": ['UE 1 – Economie & Gestion', "Finance Immobilière", "Économie Immobilière", "UE 2 – Droit", "Gestion de la Copropriété", "Droit des Sols et de la Construction", "UE 3 – Aménagement & Urbanisme", "Pathologie du Bâtiment et Suivi de Travaux", "Expertise et Evaluation Immobilière", "UE 4 – Compétences Professionnalisantes", "Immersion Professionnelle", "Atelier Urbain II", "Panorama de l'Immobilier", "Mémoire de Recherche", "Real Estate English", "ESPI Inside"],
            
            "BG_TP_1": ['UE 1 – Economie & Gestion', "Économie Générale", "Outils d'Analyse Economique", "Organisations, Stratégies et Innovations I", "Microéconomie I", "Introduction à la Finance", "Marketing & Prospection", "Mathématiques Financières", "UE 2 – Droit", "Introduction au Droit", "Droit des Contrats", "Droit des Biens", "Droit de la Copropriété I", "Droit des Baux d'Habitation", "UE 3 – Aménagement & Urbanisme", "Introduction aux Méthodes d'Analyse et de Représentation Spatiale", "Histoire Urbaine et Architecture", "UE 4 – Compétences Professionnalisantes", "Immersion Professionnelle", "Real Estate English", "Panorama de l'Immobilier", "Expression Ecrite et Orale", "Gestion du Travail", "Déontologie et Ethique Professionnelle", "ESPI Career Services", "Gestion de Projet" "ESPI Inside"],
            "BG_TP_2": ['UE 4 – Compétences Professionnalisantes', "Immersion Professionnelle", "Real Estate English & TOEFL"],
            "BG_TP_3": ['UE 1 – Economie & Gestion', "Microéconomie II", "Organisations, Stratégies et Innovations II", "Pratique de Gestion Locative I", "Transactions Résidentielles", "UE 2 – Droit", "Droit de la Vente Immobilière", "Droit de la Copropriété II", "Droit de l'Urbanisme", "Droit des Baux Commerciaux", "Droit des Baux d'Habitation", "Fiscalité Générale", "UE 3 – Aménagement & Urbanisme", "Technologie du Bâtiment", "Histoire Urbaine et Architecture", "Immobilier et Dynamiques Urbaines", "UE 4 – Compétences Professionnalisantes", "Real Estate English", "Panorama de l'Immobilier", "Expression Ecrite et Orale", "ESPI Inside"],
            
            "BG_TP_4": ["UE 4 – Compétences Professionnalisantes", "Mobilité Internationale Études"], 
            "BG_TP_5": ['UE 1 – Economie & Gestion', "Économie Urbaine", "Pratique de Gestion Locative II", "Management de Projet Immobilier", "Finance Immobilière", "UE 2 – Droit", "Droit de la Transaction Immobilière", "Droit de l'Environnement", "Fiscalité Immobilière", "Gestion de la Copropriété", "Droit des Sols et de la Construction", "UE 3 – Aménagement & Urbanisme", "Habitat et Développement Durable", "Pathologie du Bâtiment et Suivi de Travaux", "Expertise et Evaluation Immobiliere", "UE 4 – Compétences Professionnalisantes", "Real Estate English", "Panorama de l'Immobilier", "Expression Ecrite et Orale", "Méthodologie de la Recherche", "ESPI Inside", "Atelier Urbain"],
            "BG_TP_6": ["UE 4 – Compétences Professionnalisantes", "Immersion Professionnelle", "Mémoire de Recherche", "Real Estate English"],
        }
        
        progress_data[session_id] = 30  # Progression à 30%
        await update_progress(session_id, 30)
        
        column_configs = {
            "MAPI": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 23,
                'nom_site_column_index_template': 24,
                'code_groupe_column_index_template': 25,
                'nom_groupe_column_index_template': 26,
                'etendu_groupe_column_index_template': 27,
                'duree_justifie_column_index_template': 28,
                'duree_non_justifie_column_index_template': 29,
                'duree_retard_column_index_template': 30,
                'appreciation_column_index_template': 31
            },
            "MAGI": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 23,
                'nom_site_column_index_template': 24,
                'code_groupe_column_index_template': 25,
                'nom_groupe_column_index_template': 26,
                'etendu_groupe_column_index_template': 27,
                'duree_justifie_column_index_template': 28,
                'duree_non_justifie_column_index_template': 29,
                'duree_retard_column_index_template': 30,
                'appreciation_column_index_template': 31
            },
            "MEFIM": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 23,
                'nom_site_column_index_template': 24,
                'code_groupe_column_index_template': 25,
                'nom_groupe_column_index_template': 26,
                'etendu_groupe_column_index_template': 27,
                'duree_justifie_column_index_template': 28,
                'duree_non_justifie_column_index_template': 29,
                'duree_retard_column_index_template': 30,
                'appreciation_column_index_template': 31
            },
            "MAPI_S2": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 23,
                'nom_site_column_index_template': 24,
                'code_groupe_column_index_template': 25,
                'nom_groupe_column_index_template': 26,
                'etendu_groupe_column_index_template': 27,
                'duree_justifie_column_index_template': 28,
                'duree_non_justifie_column_index_template': 29,
                'duree_retard_column_index_template': 30,
                'appreciation_column_index_template': 31
            },
            "MAGI_S2": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 23,
                'nom_site_column_index_template': 24,
                'code_groupe_column_index_template': 25,
                'nom_groupe_column_index_template': 26,
                'etendu_groupe_column_index_template': 27,
                'duree_justifie_column_index_template': 28,
                'duree_non_justifie_column_index_template': 29,
                'duree_retard_column_index_template': 30,
                'appreciation_column_index_template': 31
            },
            "MEFIM_S2": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 23,
                'nom_site_column_index_template': 24,
                'code_groupe_column_index_template': 25,
                'nom_groupe_column_index_template': 26,
                'etendu_groupe_column_index_template': 27,
                'duree_justifie_column_index_template': 28,
                'duree_non_justifie_column_index_template': 29,
                'duree_retard_column_index_template': 30,
                'appreciation_column_index_template': 31
            },
            "MAPI_S3": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 21,
                'nom_site_column_index_template': 22,
                'code_groupe_column_index_template': 23,
                'nom_groupe_column_index_template': 24,
                'etendu_groupe_column_index_template': 25,
                'duree_justifie_column_index_template': 26,
                'duree_non_justifie_column_index_template': 27,
                'duree_retard_column_index_template': 28,
                'appreciation_column_index_template': 29
            },
            "MAGI_S3": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 20,
                'nom_site_column_index_template': 21,
                'code_groupe_column_index_template': 22,
                'nom_groupe_column_index_template': 23,
                'etendu_groupe_column_index_template': 24,
                'duree_justifie_column_index_template': 25,
                'duree_non_justifie_column_index_template': 26,
                'duree_retard_column_index_template': 27,
                'appreciation_column_index_template': 28
            },
            "MEFIM_S3": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 20,
                'nom_site_column_index_template': 21,
                'code_groupe_column_index_template': 22,
                'nom_groupe_column_index_template': 23,
                'etendu_groupe_column_index_template': 24,
                'duree_justifie_column_index_template': 25,
                'duree_non_justifie_column_index_template': 26,
                'duree_retard_column_index_template': 27,
                'appreciation_column_index_template': 28
            },
            "MAPI_S4": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 18,
                'nom_site_column_index_template': 19,
                'code_groupe_column_index_template': 20,
                'nom_groupe_column_index_template': 21,
                'etendu_groupe_column_index_template': 22,
                'duree_justifie_column_index_template': 23,
                'duree_non_justifie_column_index_template': 24,
                'duree_retard_column_index_template': 25,
                'appreciation_column_index_template': 26
            },
            "MAGI_S4": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 18,
                'nom_site_column_index_template': 19,
                'code_groupe_column_index_template': 20,
                'nom_groupe_column_index_template': 21,
                'etendu_groupe_column_index_template': 22,
                'duree_justifie_column_index_template': 23,
                'duree_non_justifie_column_index_template': 24,
                'duree_retard_column_index_template': 25,
                'appreciation_column_index_template': 26
            },
            "MEFIM_S4": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 18,
                'nom_site_column_index_template': 19,
                'code_groupe_column_index_template': 20,
                'nom_groupe_column_index_template': 21,
                'etendu_groupe_column_index_template': 22,
                'duree_justifie_column_index_template': 23,
                'duree_non_justifie_column_index_template': 24,
                'duree_retard_column_index_template': 25,
                'appreciation_column_index_template': 26
            },
            "BG_ALT_S1": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 21,
                'nom_site_column_index_template': 22,
                'code_groupe_column_index_template': 23,
                'nom_groupe_column_index_template': 24,
                'etendu_groupe_column_index_template': 25,
                'duree_justifie_column_index_template': 26,
                'duree_non_justifie_column_index_template': 27,
                'duree_retard_column_index_template': 28,
                'appreciation_column_index_template': 29
            },
            "BG_ALT_S2": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 22,
                'nom_site_column_index_template': 23,
                'code_groupe_column_index_template': 24,
                'nom_groupe_column_index_template': 25,
                'etendu_groupe_column_index_template': 26,
                'duree_justifie_column_index_template': 27,
                'duree_non_justifie_column_index_template': 28,
                'duree_retard_column_index_template': 29,
                'appreciation_column_index_template': 30
            }, 
            "BG_ALT_S3": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 20,
                'nom_site_column_index_template': 21,
                'code_groupe_column_index_template': 22,
                'nom_groupe_column_index_template': 23,
                'etendu_groupe_column_index_template': 24,
                'duree_justifie_column_index_template': 25,
                'duree_non_justifie_column_index_template': 26,
                'duree_retard_column_index_template': 27,
                'appreciation_column_index_template': 28
            },
            "BG_ALT_S4": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 19,
                'nom_site_column_index_template': 20,
                'code_groupe_column_index_template': 21,
                'nom_groupe_column_index_template': 22,
                'etendu_groupe_column_index_template': 23,
                'duree_justifie_column_index_template': 24,
                'duree_non_justifie_column_index_template': 25,
                'duree_retard_column_index_template': 26,
                'appreciation_column_index_template': 27
            },
            "BG_ALT_S5": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 21,
                'nom_site_column_index_template': 22,
                'code_groupe_column_index_template': 23,
                'nom_groupe_column_index_template': 24,
                'etendu_groupe_column_index_template': 25,
                'duree_justifie_column_index_template': 26,
                'duree_non_justifie_column_index_template': 27,
                'duree_retard_column_index_template': 28,
                'appreciation_column_index_template': 29
            },
            "BG_ALT_S6": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 19,
                'nom_site_column_index_template': 20,
                'code_groupe_column_index_template': 21,
                'nom_groupe_column_index_template': 22,
                'etendu_groupe_column_index_template': 23,
                'duree_justifie_column_index_template': 24,
                'duree_non_justifie_column_index_template': 25,
                'duree_retard_column_index_template': 26,
                'appreciation_column_index_template': 27
            },
            "BG_TP_S1": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 29,
                'nom_site_column_index_template': 30,
                'code_groupe_column_index_template': 31,
                'nom_groupe_column_index_template': 32,
                'etendu_groupe_column_index_template': 33,
                'duree_justifie_column_index_template': 34,
                'duree_non_justifie_column_index_template': 35,
                'duree_retard_column_index_template': 36,
                'appreciation_column_index_template': 37
            },
            "BG_TP_S2": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 6,
                'nom_site_column_index_template': 7,
                'code_groupe_column_index_template': 8,
                'nom_groupe_column_index_template': 9,
                'etendu_groupe_column_index_template': 10,
                'duree_justifie_column_index_template': 11,
                'duree_non_justifie_column_index_template': 12,
                'duree_retard_column_index_template': 13,
                'appreciation_column_index_template': 14
            },
            "BG_TP_S3": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 22,
                'nom_site_column_index_template': 23,
                'code_groupe_column_index_template': 24,
                'nom_groupe_column_index_template': 25,
                'etendu_groupe_column_index_template': 26,
                'duree_justifie_column_index_template': 27,
                'duree_non_justifie_column_index_template': 28,
                'duree_retard_column_index_template': 29,
                'appreciation_column_index_template': 30
            },
            "BG_TP_S4": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 5,
                'nom_site_column_index_template': 6,
                'code_groupe_column_index_template': 7,
                'nom_groupe_column_index_template': 8,
                'etendu_groupe_column_index_template': 9,
                'duree_justifie_column_index_template': 10,
                'duree_non_justifie_column_index_template': 11,
                'duree_retard_column_index_template': 12,
                'appreciation_column_index_template': 13
            },
            "BG_TP_S5": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 1,
                'date_naissance_column_index_template': 26,
                'nom_site_column_index_template': 27,
                'code_groupe_column_index_template': 28,
                'nom_groupe_column_index_template': 29,
                'etendu_groupe_column_index_template': 30,
                'duree_justifie_column_index_template': 31,
                'duree_non_justifie_column_index_template': 32,
                'duree_retard_column_index_template': 33,
                'appreciation_column_index_template': 34
            },
            "BG_TP_S6": {
                'name_column_index_uploaded': 2,
                'name_column_index_template': 2,
                'code_apprenant_column_index_template': 7,
                'date_naissance_column_index_template': 8,
                'nom_site_column_index_template': 9,
                'code_groupe_column_index_template': 10,
                'nom_groupe_column_index_template': 11,
                'etendu_groupe_column_index_template': 12,
                'duree_justifie_column_index_template': 13,
                'duree_non_justifie_column_index_template': 14,
                'duree_retard_column_index_template': 15,
                'appreciation_column_index_template': 16
            },
        }

        template_to_use = None
        columns_config = None
        for template, values in matching_values.items():
            if uploaded_values[:len(values)] == values:
                logger.debug(f"Matching template found: {template}")
                template_to_use = templates[template]
                columns_config = column_configs[template]
                break

        if not template_to_use or not columns_config:
            logger.error("No matching template found for the uploaded Excel data.")
            raise HTTPException(status_code=400, detail="No matching template found")

        # Log the columns to be processed
        logger.debug(f"Using template: {template_to_use}")
        logger.debug(f"Column config: {columns_config}")

        # Mise à jour des informations de groupe pour les apprenants
        for apprenant_id, frequentes_info in frequentes_dict.items():
            if apprenant_id in api_data:
                api_data[apprenant_id]['informations_frequentation'] = frequentes_info
                
                if not api_data[apprenant_id].get('informationsCourantes', {}).get('codeGroupe'):
                    api_data[apprenant_id]['informationsCourantes'] = api_data[apprenant_id].get('informationsCourantes', {})
                    api_data[apprenant_id]['informationsCourantes']['codeGroupe'] = frequentes_info.get('codeGroupe')

        # Processer le fichier et créer le fichier Excel final
        # Processer le fichier et créer le fichier Excel final
        template_wb = await process_file(uploaded_wb, template_to_use, columns_config, class_name, current_periode, previous_periode, api_data, frequentes_dict)

        progress_data[session_id] = 35  # Progression à 40%
        await update_progress(session_id, 35)

        appreciations = extract_appreciations_from_word(temp_word_path)
        logger.debug(f"Extracted appreciations: {appreciations}")

        template_wb = update_excel_with_appreciations(template_wb, appreciations, columns_config)

        template_name = os.path.basename(template_to_use).replace('.xlsx', '')
        output_path = os.path.join(settings.DOCUMENTS_DIR, f'{template_name}.xlsx')
        logger.debug(f"Saving processed workbook to {output_path}")
        template_wb.save(output_path)

        # Génération et création des bulletins PDF
        progress_data[session_id] = 50  # Progression à 50%
        await update_progress(session_id, 50)

        # Déterminer le répertoire de sortie et nettoyer les fichiers existants
        output_dir = os.path.join(settings.OUTPUT_DIR, f'bulletins-{class_name}')
        clean_output_directory(output_dir)
        os.makedirs(output_dir, exist_ok=True)
        
        bulletin_paths = process_excel_file(output_path, output_dir)

        logger.debug(f"Generated bulletins: {bulletin_paths}")
        
        convert_docx_to_pdf(output_dir)
        
        progress_data[session_id] = 70  # Progression à 55%
        await update_progress(session_id, 70)

        pdf_bulletin_paths = [
            os.path.join(output_dir, filename.replace('.docx', '.pdf'))
            for filename in os.listdir(output_dir)
            if filename.endswith('.pdf')
        ]
        
        progress_data[session_id] = 90  # Progression à 90%
        await update_progress(session_id, 90)

        # Création d'un fichier ZIP avec les PDF générés
        zip_filename = os.path.join(settings.DOWNLOAD_DIR, f'bulletins-{class_id}.zip')
        logger.debug(f"Creating ZIP file at {zip_filename}")
        with zipfile.ZipFile(zip_filename, 'w') as zipf:
            for pdf_path in pdf_bulletin_paths:
                zipf.write(pdf_path, os.path.basename(pdf_path))

        for bulletin_path in bulletin_paths:
            os.remove(bulletin_path)

        logger.debug(f"All bulletins processed and zipped successfully for class {class_name} (ID: {class_id}).")
        return JSONResponse(content={"message": f"Bulletins for {class_name} (ID: {class_id}) generated and zipped successfully", "zip_path": zip_filename})

    except Exception as e:
        logger.error("Failed to process the file and generate bulletins", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/import-bulletins-from-directory")
async def import_bulletins_from_directory():
    bulletin_dir = os.path.join(os.getenv('BASE_DIR', '/code'), 'outputs', 'bulletins')

    if not os.path.exists(bulletin_dir):
        logger.error(f"Bulletin directory not found: {bulletin_dir}")
        raise HTTPException(status_code=404, detail="Bulletin directory not found")

    import_errors = []

    for pdf_file in os.listdir(bulletin_dir):
        if pdf_file.endswith('.pdf'):
            pdf_path = os.path.join(bulletin_dir, pdf_file)

            # Log the PDF being processed
            logger.info(f"Processing PDF: {pdf_path}")

            # Extraire le code apprenant depuis le PDF (Assurez-vous d'avoir la logique pour cela)
            code_apprenant = extract_code_apprenant(pdf_path)
            if not code_apprenant:
                logger.error(f"Failed to extract code_apprenant from {pdf_path}")
                import_errors.append({
                    "file": os.path.basename(pdf_path),
                    "error": "Failed to extract code_apprenant"
                })
                continue

            logger.info(f"Extracted code_apprenant: {code_apprenant} from {pdf_path}")

            try:
                if not import_document_to_yparéo(pdf_path, code_apprenant):
                    logger.error(f"Failed to import PDF: {pdf_path}")
                    import_errors.append({
                        "file": os.path.basename(pdf_path),
                        "error": f"Failed to import {pdf_path}"
                    })
            except Exception as e:
                logger.error(f"Exception while importing PDF {pdf_path}", exc_info=True)
                import_errors.append({
                    "file": os.path.basename(pdf_path),
                    "error": str(e)
                })

    if import_errors:
        logger.warning(f"Some bulletins failed to import: {import_errors}")
        return {"message": "Some bulletins failed to import", "errors": import_errors}
    else:
        logger.info("All bulletins imported successfully")
        return {"message": "All bulletins imported successfully"}


@router.get("/download-zip/{filename}")
async def download_zip(filename: str):
    zip_path = os.path.join(settings.DOWNLOAD_DIR, filename)
    if not os.path.exists(zip_path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(path=zip_path, filename=filename, media_type='application/zip')